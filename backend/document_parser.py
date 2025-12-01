"""
Document parser for extracting content from DOCX files and converting to XML/JSON
Also supports editing DOCX files and saving changes
"""
from docx import Document as DocxDocument
from xml.etree.ElementTree import Element, SubElement, tostring
import json
import re
from typing import Dict, Any, List
import copy

def parse_docx_to_xml(file_path: str) -> str:
    """
    Parse DOCX file and convert to XML format
    """
    doc = DocxDocument(file_path)
    
    # Create root XML element
    root = Element("document")
    title_elem = SubElement(root, "title")
    title_elem.text = "Document Content"
    
    # Parse paragraphs
    content_elem = SubElement(root, "content")
    
    for para in doc.paragraphs:
        if para.text.strip():
            para_elem = SubElement(content_elem, "paragraph")
            para_elem.text = para.text
    
    # Parse tables
    for table in doc.tables:
        table_elem = SubElement(content_elem, "table")
        for row in table.rows:
            row_elem = SubElement(table_elem, "row")
            for cell in row.cells:
                cell_elem = SubElement(row_elem, "cell")
                cell_elem.text = cell.text
    
    # Convert to string
    xml_string = tostring(root, encoding='unicode')
    return xml_string

def extract_structured_data(xml_content: str) -> Dict[str, Any]:
    """
    Extract structured data from XML content
    Identifies materials, equipment, methods, etc.
    """
    from xml.etree.ElementTree import fromstring
    
    root = fromstring(xml_content)
    
    structured_data = {
        "materials": [],
        "equipment": [],
        "methods": [],
        "metadata": {}
    }
    
    # Extract paragraphs
    paragraphs = root.findall(".//paragraph")
    for para in paragraphs:
        text = para.text or ""
        
        # Look for materials
        if any(keyword in text.lower() for keyword in ["material", "reagent", "buffer", "solution"]):
            material = extract_material_info(text)
            if material:
                structured_data["materials"].append(material)
        
        # Look for equipment
        if any(keyword in text.lower() for keyword in ["equipment", "instrument", "device", "apparatus"]):
            equipment = extract_equipment_info(text)
            if equipment:
                structured_data["equipment"].append(equipment)
        
        # Look for methods
        if any(keyword in text.lower() for keyword in ["method", "procedure", "protocol", "process"]):
            structured_data["methods"].append(text)
    
    return structured_data

def extract_material_info(text: str) -> Dict[str, Any]:
    """Extract material information from text"""
    material = {
        "name": None,
        "catalog_number": None,
        "supplier": None,
        "lot_number": None
    }
    
    # Extract name (usually first part)
    name_match = re.search(r'([A-Z][a-zA-Z\s]+(?:Solution|Buffer|Reagent|Material))', text)
    if name_match:
        material["name"] = name_match.group(1).strip()
    
    # Extract catalog number
    catalog_match = re.search(r'[Cc]atalog[:\s]+([A-Z0-9\-]+)', text, re.IGNORECASE)
    if catalog_match:
        material["catalog_number"] = catalog_match.group(1)
    else:
        material["catalog_number"] = None  # Missing
    
    # Extract supplier
    supplier_match = re.search(r'[Ss]upplier[:\s]+([A-Za-z\s\-]+)', text, re.IGNORECASE)
    if supplier_match:
        material["supplier"] = supplier_match.group(1).strip()
    
    # Extract lot number
    lot_match = re.search(r'[Ll]ot[:\s]+([A-Z0-9\-]+)', text, re.IGNORECASE)
    if lot_match:
        material["lot_number"] = lot_match.group(1)
    
    return material if material["name"] else None

def extract_equipment_info(text: str) -> Dict[str, Any]:
    """Extract equipment information from text"""
    equipment = {
        "name": None,
        "configuration": None,
        "model_number": None,
        "serial_number": None
    }
    
    # Extract name
    name_match = re.search(r'([A-Z][a-zA-Z\s]+(?:Incubator|Centrifuge|Chamber|Device|System))', text)
    if name_match:
        equipment["name"] = name_match.group(1).strip()
    
    # Extract configuration
    config_match = re.search(r'[Cc]onfiguration[:\s]+([A-Za-z\s\-]+)', text, re.IGNORECASE)
    if config_match:
        config = config_match.group(1).strip()
        equipment["configuration"] = None if config.lower() == "none" else config
    else:
        equipment["configuration"] = None  # Missing
    
    # Extract model number
    model_match = re.search(r'[Mm]odel[:\s]+([A-Z0-9\-]+)', text, re.IGNORECASE)
    if model_match:
        equipment["model_number"] = model_match.group(1)
    
    # Extract serial number
    serial_match = re.search(r'[Ss]erial[:\s]+([A-Z0-9\-]+)', text, re.IGNORECASE)
    if serial_match:
        equipment["serial_number"] = serial_match.group(1)
    
    return equipment if equipment["name"] else None

def identify_gaps(structured_data: Dict[str, Any], document_structure: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Identify gaps in the structured data and table cells
    Returns a dictionary with gap information
    """
    gaps = {
        "materials": [],
        "equipment": [],
        "table_cells": [],
        "total_gaps": 0
    }
    
    # Check materials for missing catalog numbers
    for idx, material in enumerate(structured_data.get("materials", [])):
        if not material.get("catalog_number"):
            gaps["materials"].append({
                "index": idx,
                "field": "catalog_number",
                "material_name": material.get("name", "Unknown"),
                "status": "missing"
            })
            gaps["total_gaps"] += 1
    
    # Check equipment for missing/invalid configurations
    for idx, equipment in enumerate(structured_data.get("equipment", [])):
        if not equipment.get("configuration") or equipment.get("configuration") == "None":
            gaps["equipment"].append({
                "index": idx,
                "field": "configuration",
                "equipment_name": equipment.get("name", "Unknown"),
                "status": "invalid"
            })
            gaps["total_gaps"] += 1
    
    # Check table cells for "(Pending)" or missing values
    if document_structure and document_structure.get("tables"):
        for table_idx, table in enumerate(document_structure.get("tables", [])):
            # Get table name and column headers
            table_name = table.get("name") or table.get("id", f"Table {table_idx + 1}")
            column_headers = table.get("column_headers", [])
            
            for row_idx, row in enumerate(table.get("rows", [])):
                for cell_idx, cell in enumerate(row.get("cells", [])):
                    cell_text = cell.get("text", "").strip()
                    cell_text_lower = cell_text.lower() if cell_text else ""
                    
                    # Check for pending values - both "(Pending)" and standalone "Pending" word
                    # Check for various formats: "(Pending)", "(pending)", "(PENDING)", "Pending)", "(Pending", "Pending", etc.
                    is_pending = (
                        "(Pending)" in cell_text or 
                        "(pending)" in cell_text_lower or 
                        "(PENDING)" in cell_text or
                        "Pending)" in cell_text or 
                        "(Pending" in cell_text or
                        "pending)" in cell_text_lower or
                        "(pending" in cell_text_lower or
                        "pending" in cell_text_lower or 
                        "Pending" in cell_text or 
                        "PENDING" in cell_text
                    )
                    
                    # Check for null values - exact match or contains "null" keyword
                    null_keywords = ["null", "none", "nil", "n/a", "na", "n.a."]
                    # Check if exact match or contains "null" as a word
                    is_null = (cell_text_lower in null_keywords or 
                              "null" in cell_text_lower or
                              (cell_text_lower.strip() and cell_text_lower.strip() in null_keywords))
                    
                    # Check for missing values
                    missing_keywords = ["missing", "not provided", "not available", "unavailable", 
                                       "tbd", "to be determined", "t.b.d.", "tba", "to be announced",
                                       "unknown", "unk"]
                    is_missing_text = cell_text_lower in missing_keywords or any(keyword in cell_text_lower for keyword in missing_keywords)
                    
                    # Check for empty or whitespace-only cells
                    is_empty = not cell_text or len(cell_text.strip()) == 0
                    
                    # Check if cell contains null/missing keywords (for partial matches)
                    contains_null_keyword = any(keyword in cell_text_lower for keyword in 
                                               ["null", "missing", "not provided", "not available", "unavailable"])
                    
                    # Determine if this is a gap (only count once per cell)
                    is_gap = is_pending or is_null or is_missing_text or is_empty or contains_null_keyword
                    
                    if is_gap:
                        # Determine gap type (prioritize pending, then null, then missing, then empty)
                        if is_pending:
                            gap_type = "Pending"
                            gap_status = "pending"
                        elif is_null or "null" in cell_text_lower:
                            gap_type = "Null value"
                            gap_status = "null"
                        elif is_missing_text or contains_null_keyword:
                            gap_type = "Missing value"
                            gap_status = "missing"
                        elif is_empty:
                            gap_type = "Empty cell"
                            gap_status = "empty"
                        else:
                            gap_type = "Data quality gap"
                            gap_status = "gap"
                        
                        # Get column header name if available, otherwise use column number
                        field_name = column_headers[cell_idx] if cell_idx < len(column_headers) and column_headers[cell_idx] else f"Column {cell_idx + 1}"
                        
                        # Build meaningful description with table name and field name
                        if row_idx == 0 and len(column_headers) > 0:
                            # If this is the header row, don't include row number
                            description = f"{table_name} - Field: {field_name} - {gap_type}"
                        else:
                            description = f"{table_name} - Row {row_idx + 1}, Field: {field_name} - {gap_type}"
                        
                        gaps["table_cells"].append({
                            "table_id": table.get("id", f"table_{table_idx}"),
                            "table_name": table_name,
                            "field_name": field_name,
                            "row": row_idx,
                            "col": cell_idx,
                            "text": cell_text or "(empty)",
                            "issue": gap_type,
                            "status": gap_status,
                            "description": description
                        })
                        gaps["total_gaps"] += 1
    
    # Recalculate total_gaps to ensure accuracy (sum of all gap types)
    gaps["total_gaps"] = (
        len(gaps["materials"]) + 
        len(gaps["equipment"]) + 
        len(gaps["table_cells"])
    )
    
    return gaps

def get_docx_structure(file_path: str, preserved_table_metadata: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Extract structured content from DOCX file including tables
    Returns paragraphs and tables as structured data
    
    Args:
        file_path: Path to the DOCX file
        preserved_table_metadata: Optional dict mapping table indices to preserved metadata (name, id, column_headers)
    """
    doc = DocxDocument(file_path)
    
    # Prepare preserved metadata lookup
    preserved_metadata = preserved_table_metadata or {}
    
    structure = {
        "paragraphs": [],
        "tables": []
    }
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            structure["paragraphs"].append({
                "text": para.text,
                "style": para.style.name if para.style else None
            })
    
    # Extract tables with table names and column headers
    # First, create a mapping of all document elements to track table positions
    all_elements = []
    table_idx_map = {}
    
    # Collect all paragraphs and tables in order
    para_idx = 0
    for element in doc.element.body:
        if element.tag.endswith('}p'):  # Paragraph
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break
            if para and para.text.strip():
                all_elements.append(('paragraph', para_idx, para.text.strip()))
                para_idx += 1
        elif element.tag.endswith('}tbl'):  # Table
            tbl_idx = None
            for idx, tbl in enumerate(doc.tables):
                if tbl._element == element:
                    tbl_idx = idx
                    break
            if tbl_idx is not None:
                all_elements.append(('table', tbl_idx, None))
                table_idx_map[tbl_idx] = len(all_elements) - 1
    
    # First pass: Collect all potential table names to detect if we're using headers incorrectly
    # If multiple tables have short "names" that look like headers, we know we're extracting wrong
    potential_table_names = []
    
    for table_idx, table in enumerate(doc.tables):
        table_data = {
            "id": f"table_{table_idx}",
            "name": None,  # Will be set from paragraph before table or first row
            "column_headers": [],  # Will be set from first row
            "rows": []
        }
        
        # CRITICAL: Check preserved metadata FIRST before any extraction
        # If preserved metadata exists, use it directly and skip extraction
        was_preserved = False
        table_name = None
        column_headers = []
        if preserved_metadata and table_idx in preserved_metadata:
            preserved = preserved_metadata[table_idx]
            if preserved and preserved.get("name"):
                table_name = preserved["name"]
                was_preserved = True
                print(f"Using preserved table name for table {table_idx}: {table_name}")
            if preserved and preserved.get("id"):
                table_data["id"] = preserved["id"]
            if preserved and preserved.get("column_headers"):
                column_headers = preserved["column_headers"]
        
        # Only run extraction logic if we don't have preserved metadata
        table_name_from_paragraph = False  # Track if we got name from paragraph
        if not was_preserved:
            # Initialize column_headers for extraction
            column_headers = []
            # Try to get table name from paragraph before this table
            table_name = None
        if table_idx in table_idx_map:
            table_position = table_idx_map[table_idx]
            # Look for paragraph before this table (check up to 10 paragraphs back to find table titles)
            # Some documents might have multiple paragraphs between sections
            for i in range(table_position - 1, max(-1, table_position - 11), -1):
                if all_elements[i][0] == 'paragraph':
                    para_text = all_elements[i][2]
                    # Check if paragraph looks like a table title (short, no punctuation at end, title case)
                    if para_text and len(para_text) < 100 and not para_text.endswith('.'):
                        para_lower = para_text.lower().strip()
                        para_text_stripped = para_text.strip()
                        
                        # REJECT if it looks like a column header (contains header patterns)
                        has_header_patterns = any(pattern in para_lower for pattern in 
                                                  ['_', '#', 'atn', 'attn', 'ctg', 'bg_', 'bg ', 'material_', 'material ', 'bg_atn', 'bg_attn', 'material_ctg']) or \
                                             (len(para_text_stripped) < 20 and ('_' in para_text_stripped or '#' in para_text_stripped)) or \
                                             para_lower in ['bg_atn', 'bg_attn', 'bg attn', 'material_ctg', 'material_ctg #', 'material ctg']
                        
                        if has_header_patterns:
                            continue  # Skip this paragraph, it's likely not a table name
                        
                        # Common table name patterns - these are good indicators
                        # Check for specific known table names like "Equipment Configuration", "Verification Activities"
                        if any(keyword in para_lower for keyword in 
                              ['table', 'materials', 'equipment', 'method', 'procedure', 'result', 
                               'data', 'summary', 'list', 'specification', 'requirement', 'verification', 'activities',
                               'configuration', 'test', 'analysis', 'activity', 'activities']):
                            table_name = para_text_stripped
                            table_name_from_paragraph = True
                            print(f"Found table name from paragraph for table {table_idx}: {table_name}")
                            break
                        # Or if it's a short line (likely a title) - but only if it doesn't look like a header
                        elif len(para_text_stripped) < 60 and len(para_text_stripped) > 5:
                            # Additional check: must contain words (not just acronyms)
                            # Table names usually have multiple words or are descriptive
                            word_count = len(para_text_stripped.split())
                            # Accept if it has 2+ words, or if it's a single descriptive word (like "Equipment", "Materials")
                            if word_count >= 2 or (word_count == 1 and len(para_text_stripped) > 8):
                                # Make sure it's not a known header
                                if para_lower not in ['bg_atn', 'bg_attn', 'material_ctg']:
                                    table_name = para_text_stripped
                                    table_name_from_paragraph = True
                                    print(f"Found table name from paragraph (short) for table {table_idx}: {table_name}")
                                    break
        
        # Extract column headers from first row (only if we don't have preserved metadata)
        if not was_preserved and len(table.rows) > 0:
            first_row = table.rows[0]
            
            # Check if first row might be a table name row instead of headers
            first_row_texts = [cell.text.strip() for cell in first_row.cells]
            is_title_row = False
            
            # Heuristics to determine if first row is a title row:
            # - It should NOT look like column headers (short values, common header patterns)
            # - Column headers often contain: underscores, #, short acronyms, common words like "name", "id", "date"
            header_indicators = ['_', '#', 'id', 'name', 'date', 'type', 'status', 'code', 'num', 'qty', 'qty', 'atn', 'ctg']
            looks_like_headers = any(
                any(indicator in text.lower() for indicator in header_indicators) or 
                (len(text) < 15 and ('_' in text or text.isupper() or '#' in text))
                for text in first_row_texts
            )
            
            # CRITICAL: We should NEVER use first row as table name if it contains header patterns
            # Column headers like "BG_ATN", "Material_CTG #" should NEVER be used as table names
            # Only consider first row as table name if:
            # 1. We don't have a table name from paragraph AND
            # 2. First row doesn't look like headers AND  
            # 3. It's a single merged cell with substantial text AND
            # 4. It doesn't contain ANY header-like patterns
            
            # Check for header patterns more strictly
            # Include specific patterns like "BG_ATTN", "BG_ATN", "Material_CTG #", etc.
            first_row_has_header_patterns = any(
                any(pattern in text.lower() for pattern in ['_', '#', 'atn', 'attn', 'ctg', 'bg', 'material_', 'material ']) or
                (len(text) < 20 and ('_' in text or '#' in text)) or
                'bg_' in text.lower() or 'bg ' in text.lower() or
                'material_ctg' in text.lower() or 'material ctg' in text.lower()
                for text in first_row_texts
            )
            
            # If first row has merged cells or all cells have same text, might be title
            # BUT only if it doesn't look like headers AND we don't have a name from paragraph
            # AND it doesn't contain header patterns
            if not table_name and not first_row_has_header_patterns:  # Only consider if we don't have a name from paragraph
                if len(first_row.cells) == 1 and not looks_like_headers:
                    # Single cell spanning all columns - likely table title (if not a header)
                    if first_row_texts[0] and len(first_row_texts[0]) > 20:  # Increased minimum length
                        # Additional validation: make sure it doesn't contain header patterns
                        cell_text_lower = first_row_texts[0].lower()
                        known_headers_in_cell = any(header in cell_text_lower for header in ['bg_atn', 'bg_attn', 'material_ctg'])
                        has_header_pattern = any(pattern in cell_text_lower for pattern in ['_', '#', 'atn', 'attn', 'ctg', 'bg', 'material_'])
                        if not known_headers_in_cell and not has_header_pattern:
                            table_name = first_row_texts[0]
                            is_title_row = True
                elif len(set(first_row_texts)) == 1 and len(first_row_texts[0]) > 25 and not looks_like_headers:
                    # All cells have same long text - might be title (if not headers)
                    # Additional validation
                    cell_text_lower = first_row_texts[0].lower()
                    known_headers_in_cell = any(header in cell_text_lower for header in ['bg_atn', 'bg_attn', 'material_ctg'])
                    has_header_pattern = any(pattern in cell_text_lower for pattern in ['_', '#', 'atn', 'attn', 'ctg', 'bg', 'material_'])
                    if not known_headers_in_cell and not has_header_pattern:
                        table_name = first_row_texts[0]
                        is_title_row = True
            
            # Extract headers (skip first row if it's a title)
            if is_title_row and len(table.rows) > 1:
                # Use second row as headers
                header_row = table.rows[1]
                for cell in header_row.cells:
                    header_text = cell.text.strip()
                    column_headers.append(header_text if header_text else f"Column {len(column_headers) + 1}")
            else:
                # IMPORTANT: We should NEVER use first row as table name if it looks like headers
                # This prevents column headers (like "BG_ATN", "Material_CTG #") from being used as table names
                
                # Check if first row definitely looks like headers
                # Strong indicators: contains underscores, #, short text, uppercase, common header words
                header_indicators_list = ['_', '#', 'id', 'name', 'date', 'type', 'status', 'code', 'num', 'qty', 'atn', 'ctg', 'bg', 'material']
                looks_like_headers_definite = any(
                    any(indicator in text.lower() for indicator in header_indicators_list) or
                    (len(text) < 25 and ('_' in text or text.isupper() or '#' in text)) or
                    (len(text) < 15)  # Very short text is likely a header
                    for text in first_row_texts
                )
                
                if looks_like_headers_definite:
                    # First row is DEFINITELY headers - use as column headers, never as table name
                    # DO NOT set table_name from first row if it looks like headers
                    for cell in first_row.cells:
                        header_text = cell.text.strip()
                        column_headers.append(header_text if header_text else f"Column {len(column_headers) + 1}")
                elif table_name:
                    # We have a table name from paragraph, so first row is headers
                    for cell in first_row.cells:
                        header_text = cell.text.strip()
                        column_headers.append(header_text if header_text else f"Column {len(column_headers) + 1}")
                else:
                    # No table name yet and first row doesn't look like headers
                    # Only consider first row as table name if it's a single merged cell with substantial text
                    if len(first_row.cells) == 1 and first_row_texts[0] and len(first_row_texts[0]) > 20:
                        # Single merged cell with substantial text - might be table name
                        # But only if it doesn't contain header patterns
                        text_lower = first_row_texts[0].lower()
                        known_headers_in_text = any(header in text_lower for header in ['bg_atn', 'bg_attn', 'material_ctg'])
                        has_header_pattern = any(indicator in text_lower for indicator in ['_', '#', 'atn', 'attn', 'ctg', 'bg', 'material_'])
                        if not known_headers_in_text and not has_header_pattern:
                            table_name = first_row_texts[0]
                            is_title_row = True
                            # Use second row as headers if available
                            if len(table.rows) > 1:
                                header_row = table.rows[1]
                                for cell in header_row.cells:
                                    header_text = cell.text.strip()
                                    column_headers.append(header_text if header_text else f"Column {len(column_headers) + 1}")
                        else:
                            # Contains header patterns, treat as headers
                            for cell in first_row.cells:
                                header_text = cell.text.strip()
                                column_headers.append(header_text if header_text else f"Column {len(column_headers) + 1}")
                    else:
                        # Default: treat first row as headers (safer assumption)
                        for cell in first_row.cells:
                            header_text = cell.text.strip()
                            column_headers.append(header_text if header_text else f"Column {len(column_headers) + 1}")
        
        # CRITICAL: Final validation - reject any table name that looks like a column header
        # (Only validate if we didn't preserve metadata - preserved names are trusted)
        # Column headers like "BG_ATN", "BG_ATTN", "Material_CTG #" should NEVER be used as table names
        if table_name and not was_preserved:
            table_name_lower = table_name.lower().strip()
            table_name_upper = table_name.upper().strip()
            
            # Strong header patterns that should NEVER be table names
            header_patterns = ['_', '#', 'atn', 'attn', 'ctg', 'bg_', 'bg ', 'material_ctg', 'material_', 
                              'material ', 'bg_atn', 'bg attn', 'bg_attn']
            
            # Specific known column header patterns to reject
            known_headers = ['bg_atn', 'bg_attn', 'bg attn', 'bg atn', 'material_ctg', 'material_ctg #', 
                            'material ctg', 'material ctg #', 'material_ctg#']
            
            # Check if it matches known headers exactly (case insensitive)
            matches_known_header = any(known_header.lower() in table_name_lower or table_name_lower == known_header.lower() 
                                      for known_header in known_headers)
            # Also check exact matches
            if table_name_lower in [h.lower() for h in known_headers] or \
               table_name_upper in [h.upper() for h in known_headers]:
                matches_known_header = True
                
            has_header_pattern = any(pattern in table_name_lower for pattern in header_patterns)
            is_short_with_special = len(table_name) < 25 and ('_' in table_name or '#' in table_name)
            is_very_short = len(table_name) < 15
            has_single_word = len(table_name.split()) == 1 and len(table_name) < 20
            
            # Reject if it matches known headers or looks like a header
            if matches_known_header or has_header_pattern or is_short_with_special or \
               (is_very_short and not table_name_from_paragraph) or (has_single_word and has_header_pattern):
                print(f"REJECTING: Table name '{table_name}' looks like a column header, resetting to default for table {table_idx}")
                table_name = None
                
                # Try one more time to find a better name from paragraphs if we rejected a bad one
                if not table_name_from_paragraph and table_idx in table_idx_map:
                    table_position = table_idx_map[table_idx]
                    # Look further back (up to 15 paragraphs) for a better table name
                    for i in range(table_position - 1, max(-1, table_position - 16), -1):
                        if all_elements[i][0] == 'paragraph':
                            para_text = all_elements[i][2]
                            if para_text and len(para_text) < 100 and not para_text.endswith('.'):
                                para_lower = para_text.lower().strip()
                                para_stripped = para_text.strip()
                                # Check if this paragraph looks like a table name (has keywords, not a header)
                                if any(keyword in para_lower for keyword in ['equipment', 'verification', 'materials', 'configuration', 'activities', 'method', 'procedure']) and \
                                   not any(pattern in para_lower for pattern in ['_', '#', 'atn', 'attn', 'ctg', 'bg_']):
                                    table_name = para_stripped
                                    table_name_from_paragraph = True
                                    print(f"Found alternative table name from paragraph for table {table_idx}: {table_name}")
                                    break
        
        # Set default table name if not found
        if not table_name:
            table_name = f"Table {table_idx + 1}"
        
        table_data["name"] = table_name
        table_data["column_headers"] = column_headers
        
        for row_idx, row in enumerate(table.rows):
            row_data = {
                "id": f"row_{table_idx}_{row_idx}",
                "cells": []
            }
            
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                cell_text_lower = cell_text.lower() if cell_text else ""
                
                # Check for pending values - both "(Pending)" and standalone "Pending" word
                # Check for various formats: "(Pending)", "(pending)", "(PENDING)", "Pending)", "(Pending", "Pending", etc.
                is_pending = (
                    "(Pending)" in cell_text or 
                    "(pending)" in cell_text_lower or 
                    "(PENDING)" in cell_text or
                    "Pending)" in cell_text or 
                    "(Pending" in cell_text or
                    "pending)" in cell_text_lower or
                    "(pending" in cell_text_lower or
                    "pending" in cell_text_lower or 
                    "Pending" in cell_text or 
                    "PENDING" in cell_text
                )
                
                # Check for null values - exact match or contains "null" keyword
                null_keywords = ["null", "none", "nil", "n/a", "na", "n.a."]
                # Check if exact match or contains "null" as a word
                is_null = (cell_text_lower in null_keywords or 
                          "null" in cell_text_lower or
                          (cell_text_lower.strip() and cell_text_lower.strip() in null_keywords))
                
                # Check for missing values
                missing_keywords = ["missing", "not provided", "not available", "unavailable", 
                                   "tbd", "to be determined", "t.b.d.", "tba", "to be announced",
                                   "unknown", "unk"]
                is_missing_text = (cell_text_lower in missing_keywords or 
                                 any(keyword in cell_text_lower for keyword in missing_keywords))
                
                # Check for empty or whitespace-only cells
                is_empty = not cell_text or len(cell_text.strip()) == 0
                
                # Check if cell contains null/missing keywords (for partial matches)
                contains_null_keyword = any(keyword in cell_text_lower for keyword in 
                                           ["null", "missing", "not provided", "not available", "unavailable"])
                
                # Determine if this is a gap
                has_gap = is_pending or is_null or is_missing_text or is_empty or contains_null_keyword
                
                cell_data = {
                    "id": f"cell_{table_idx}_{row_idx}_{cell_idx}",
                    "text": cell_text,
                    "row": row_idx,
                    "col": cell_idx,
                    "is_pending": is_pending,
                    "is_null": is_null or "null" in cell_text_lower,
                    "is_missing": is_missing_text or contains_null_keyword,
                    "is_empty": is_empty,
                    "has_gap": has_gap
                }
                row_data["cells"].append(cell_data)
            
            table_data["rows"].append(row_data)
        
        structure["tables"].append(table_data)
    
    return structure

def get_docx_text_content(file_path: str) -> str:
    """
    Extract text content from DOCX file for editing
    Returns formatted text that can be edited
    """
    doc = DocxDocument(file_path)
    content_lines = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            content_lines.append(para.text)
    
    for table in doc.tables:
        content_lines.append("")  # Empty line before table
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            content_lines.append(row_text)
        content_lines.append("")  # Empty line after table
    
    return "\n".join(content_lines)

def save_structure_to_docx(file_path: str, structure: Dict[str, Any]) -> str:
    """
    Save structured content (paragraphs and tables) back to DOCX file
    Preserves original document structure and formatting
    """
    # Load the existing document to preserve formatting
    try:
        doc = DocxDocument(file_path)
        
        # Update paragraphs - match by index and update text only
        # CRITICAL: Preserve paragraphs that look like section headings or table titles
        # to prevent table title paragraphs (like "3.Equipment Configuration") from being overwritten
        para_idx = 0
        for para in doc.paragraphs:
            if para.text.strip() and para_idx < len(structure.get("paragraphs", [])):
                para_data = structure["paragraphs"][para_idx]
                original_text = para.text.strip()
                new_text = para_data.get("text", "").strip()
                
                # Check if this paragraph looks like a section heading or table title
                # These should be preserved and not overwritten with incorrect data
                original_lower = original_text.lower()
                is_section_heading = (
                    # Contains section numbers like "3.", "4.", etc.
                    (len(original_text) > 0 and original_text[0].isdigit() and '.' in original_text[:5]) or
                    # Contains keywords that indicate section headings or table titles
                    any(keyword in original_lower for keyword in [
                        'equipment configuration', 'verification activities', 'materials used',
                        'equipment', 'configuration', 'verification', 'activities', 'materials',
                        'table', 'section', 'subsection'
                    ]) and len(original_text) > 10  # Must be substantial text
                )
                
                # Check if new_text looks like a column header identifier (should not replace section headings)
                new_text_lower = new_text.lower() if new_text else ""
                looks_like_header_identifier = (
                    new_text and (
                        any(pattern in new_text_lower for pattern in ['bg_atn', 'bg_attn', 'material_ctg', 'material_ctg#']) or
                        (len(new_text) < 20 and ('_' in new_text or '#' in new_text)) or
                        new_text_lower in ['bg_atn', 'bg_attn', 'material_ctg', 'material_ctg #', 'material ctg']
                    )
                )
                
                # CRITICAL: If original is a section heading, be very conservative about updating it
                # Only update if the new text is also a valid section heading or is clearly different but legitimate
                if is_section_heading:
                    # Check if new text is also a section heading (similar structure)
                    new_is_section_heading = (
                        (len(new_text) > 0 and new_text[0].isdigit() and '.' in new_text[:5]) or
                        any(keyword in new_text_lower for keyword in [
                            'equipment configuration', 'verification activities', 'materials used',
                            'equipment', 'configuration', 'verification', 'activities', 'materials'
                        ]) and len(new_text) > 10
                    )
                    
                    # If new text looks like a header identifier, definitely preserve original
                    if looks_like_header_identifier:
                        print(f"PRESERVING section heading paragraph {para_idx}: '{original_text}' (rejecting header identifier '{new_text}')")
                        para_idx += 1
                        continue  # Skip updating this paragraph
                    
                    # If new text is not a section heading and is very different, be cautious
                    # Only update if new text is substantially similar or is clearly a valid replacement
                    if not new_is_section_heading and original_text != new_text:
                        # If the new text is much shorter or looks wrong, preserve original
                        if len(new_text) < len(original_text) * 0.5 or looks_like_header_identifier:
                            print(f"PRESERVING section heading paragraph {para_idx}: '{original_text}' (new text '{new_text}' seems incorrect)")
                            para_idx += 1
                            continue
                
                # Otherwise, update normally
                if original_text != new_text:
                    # Clear existing runs and add new text (preserves paragraph style)
                    para.clear()
                    para.add_run(new_text)
                    print(f"Updated paragraph {para_idx}: '{original_text}' -> '{new_text}'")
                para_idx += 1
        
        # Update tables - match by index and update cell text only
        table_idx = 0
        for doc_table in doc.tables:
            if table_idx < len(structure.get("tables", [])):
                table_data = structure["tables"][table_idx]
                rows = table_data.get("rows", [])
                column_headers = table_data.get("column_headers", [])
                
                # Update each cell in the existing table
                for row_idx, row_data in enumerate(rows):
                    if row_idx < len(doc_table.rows):
                        cells = row_data.get("cells", [])
                        
                        # CRITICAL: For header row (row 0), use preserved column_headers instead of cell text
                        # This prevents internal identifiers like "BG_ATN" from overwriting display names like "Equipment Configuration"
                        if row_idx == 0 and column_headers:
                            # Restore header row from preserved column_headers
                            for cell_idx in range(min(len(column_headers), len(doc_table.rows[row_idx].cells))):
                                cell = doc_table.rows[row_idx].cells[cell_idx]
                                old_text = cell.text
                                new_text = column_headers[cell_idx] if cell_idx < len(column_headers) else ""
                                if new_text:  # Only update if we have a header value
                                    cell.text = new_text
                                    if old_text != new_text:
                                        print(f"Restored header table {table_idx}, cell {cell_idx}: '{old_text}' -> '{new_text}'")
                        else:
                            # For data rows (row_idx > 0), update normally
                            for cell_idx, cell_data in enumerate(cells):
                                if cell_idx < len(doc_table.rows[row_idx].cells):
                                    # Update cell text while preserving formatting
                                    cell = doc_table.rows[row_idx].cells[cell_idx]
                                    old_text = cell.text
                                    new_text = cell_data.get("text", "")
                                    cell.text = new_text
                                    if old_text != new_text:
                                        print(f"Updated table {table_idx}, row {row_idx}, cell {cell_idx}: '{old_text}' -> '{new_text}'")
                table_idx += 1
        
        # Save the updated document
        doc.save(file_path)
        return file_path
        
    except Exception as e:
        # Fallback: Create new document if update fails
        print(f"Warning: Could not update existing document, creating new one: {str(e)}")
        doc = DocxDocument()
        
        # Add paragraphs
        for para_data in structure.get("paragraphs", []):
            para = doc.add_paragraph(para_data.get("text", ""))
            if para_data.get("style"):
                try:
                    para.style = para_data["style"]
                except:
                    pass  # Style not found, use default
        
        # Add tables
        for table_data in structure.get("tables", []):
            rows = table_data.get("rows", [])
            if not rows:
                continue
            
            column_headers = table_data.get("column_headers", [])
            
            # Determine table dimensions
            num_cols = max([len(row.get("cells", [])) for row in rows], default=1)
            if column_headers:
                num_cols = max(num_cols, len(column_headers))
            num_rows = len(rows)
            
            # Create table
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Light Grid Accent 1'  # Apply a table style
            
            # Populate table cells
            for row_idx, row_data in enumerate(rows):
                cells = row_data.get("cells", [])
                
                # CRITICAL: For header row (row 0), use preserved column_headers instead of cell text
                if row_idx == 0 and column_headers:
                    # Use column_headers for header row
                    for cell_idx in range(min(len(column_headers), num_cols)):
                        cell = table.rows[row_idx].cells[cell_idx]
                        cell.text = column_headers[cell_idx] if cell_idx < len(column_headers) else ""
                else:
                    # For data rows, use cell text normally
                    for cell_idx, cell_data in enumerate(cells):
                        if row_idx < num_rows and cell_idx < num_cols:
                            cell = table.rows[row_idx].cells[cell_idx]
                            cell.text = cell_data.get("text", "")
        
        # Save the document
        doc.save(file_path)
        return file_path

def save_text_to_docx(file_path: str, text_content: str) -> str:
    """
    Save text content back to DOCX file
    Creates a new DOCX with the updated content
    """
    doc = DocxDocument()
    
    lines = text_content.split('\n')
    current_paragraph = None
    
    for line in lines:
        line = line.strip()
        if not line:
            # Empty line - add paragraph break
            if current_paragraph:
                current_paragraph = None
            continue
        
        # Check if line looks like a table row (contains |)
        if '|' in line:
            # This is a table row - for now, add as paragraph
            # In a full implementation, you'd parse and create actual tables
            para = doc.add_paragraph(line.replace('|', ' | '))
        else:
            # Regular paragraph
            para = doc.add_paragraph(line)
            current_paragraph = para
    
    # Save the document
    doc.save(file_path)
    return file_path

def get_docx_for_editing(file_path: str, preserved_table_metadata: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Get DOCX content in a format suitable for editing
    Returns both text content, structured data, and document structure (tables)
    
    Args:
        file_path: Path to the DOCX file
        preserved_table_metadata: Optional dict mapping table indices to preserved metadata (name, id, column_headers)
    """
    text_content = get_docx_text_content(file_path)
    document_structure = get_docx_structure(file_path, preserved_table_metadata)
    xml_content = parse_docx_to_xml(file_path)
    structured_data = extract_structured_data(xml_content)
    gaps = identify_gaps(structured_data, document_structure)
    
    return {
        "text_content": text_content,
        "document_structure": document_structure,  # Includes tables
        "structured_data": structured_data,
        "gaps": gaps
    }

